from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D9D9D9')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    add_bottom_border(p)


def add_body(doc, text, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_after=Pt(3)):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.63)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.6)
section.bottom_margin = Cm(1.6)
section.left_margin = Cm(1.6)
section.right_margin = Cm(1.6)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

name = doc.add_paragraph()
name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
name.paragraph_format.space_after = Pt(2)
run = name.add_run('Mark Healy')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(19)

add_body(doc, 'Cork, Ireland | +353 868915769 | healymark251@hotmail.com', space_after=Pt(6))
add_body(doc, 'CQV / Validation-Focused GMP Biotechnician', bold=True, space_after=Pt(8))

add_heading(doc, 'PROFESSIONAL PROFILE')
add_body(
    doc,
    'Validation-focused GMP professional with 5+ years of biopharma experience across downstream purification, '
    'aseptic fill-finish, QC and solid dose manufacturing. First Class Honours PGDip in Commissioning, Qualification '
    'and Validation (CQV) with formal training in URS, DQ, FAT/SAT, IQ/OQ/PQ, FMEA, CSV, P&IDs, cleanroom/facility '
    'design and tech transfer. Strong background in GDP/ALCOA+ documentation, deviations/CAPA/change control, PPQ '
    'support, PAS-X, DeltaV and TrackWise. Selected as 1 of 4 Operations representatives to support cross-functional '
    'FDA-response / validation activities on aseptic installation improvements. Open to CQV, C&Q and Validation roles '
    'in Ireland or abroad.',
    space_after=Pt(8)
)

add_heading(doc, 'CORE SKILLS')
add_body(doc, 'Validation / CQV: URS, DQ, FAT/SAT, IQ/OQ/PQ, FMEA, CSV, PPQ support, validation documentation, change control, deviation / CAPA support, FDA-response remediation support')
add_body(doc, 'Quality / Compliance: cGMP, GDP / ALCOA+, Annex 1 contamination-control mindset, audit / inspection readiness, line clearance, IPC testing, evidence capture')
add_body(doc, 'Systems: PAS-X MES (eBR / eLogs), DeltaV, TrackWise, SAP, SIMCA, LIMS, Excel, Power BI')
add_body(doc, 'Process / Equipment: Downstream purification, chromatography, UF/DF (TFF), viral filtration, CIP / SIP, Bausch + Strobel final fill, RABS / isolator, VHP readiness, parts washer / autoclave, sterile component staging / transfer, solid dose compression / blending', space_after=Pt(8))

add_heading(doc, 'PROFESSIONAL EXPERIENCE')
add_body(doc, 'BioMarin, Shanbally, Cork - Biotechnician (Downstream Purification + Final Fill) | Full time Aug 2024 - Present', bold=True)
biomarin_bullets = [
    'Cross-trained GMP biotechnician supporting Downstream Purification (Drug Substance) and Final Fill (Drug Product), executing operations and documentation to cGMP, GDP and ALCOA+ standards.',
    'Execute downstream operations including depth filtration, IMAC / CEX / HIC chromatography, pH adjust / hold (viral inactivation), UF/DF (TFF), viral filtration, DNA removal, formulation support, buffer preparation / transfers and single-use assemblies.',
    'Support PPQ execution through disciplined CPP / IPC data capture and traceable evidence in PAS-X eBR.',
    'Perform IPC sampling / testing (for example pH and conductivity) with contemporaneous documentation; escalate deviations / trends and support investigations and CAPA with QA and Engineering.',
    'Act as PAS-X MES SME for eBR / eLogs execution and troubleshooting; support DeltaV monitoring / execution and SAP transactions to prevent delays and maintain continuity.',
    'Support validated aseptic fill-finish on Bausch + Strobel equipment, including critical gowning, line support / clearance, IPC checks, sterile component staging / transfer, RABS / isolator operations and VHP readiness documentation.',
    'Selected as 1 of 4 Operations representatives to support cross-functional validation / remediation activities in response to FDA feedback on isolator installation practices.',
    'Contribute to review and improvement of aseptic installation steps to reduce operator reach into the isolator, including assessment of tooling / access changes, protective covers for critical components and procedural improvements through site change control.',
    'Support troubleshooting of needle installation and other critical setup steps to improve repeatability, access and contamination-control robustness.',
    'Area 5S representative for Downstream; completed a 5S project using standardised layout, labelling and visual controls to improve flow, reduce retrieval time and strengthen audit readiness.',
]
for b in biomarin_bullets:
    add_bullet(doc, b)

add_body(doc, 'Career Break - Travel & Personal Development (Australia) Sep 2023 - Jun 2024', bold=True, space_after=Pt(4))

add_body(doc, 'AbbVie, Carrigtwohill, Cork - Production Operator (Solid Dose) Mar 2022 - Aug 2023', bold=True)
abbvie_bullets = [
    'Operated tablet compression and blending equipment in a cGMP solid dose facility, including set-up, start-up, in-process adjustments within validated ranges, shutdown and cleaning per SOPs.',
    'Completed product changeovers on the Fette press, including tooling changes, strip-down, cleaning, reassembly and post-clean checks; documented line clearance and cleaning records.',
    'Executed batch records, logbooks and line clearance documentation to GDP / ALCOA+ standards and maintained strong right-first-time performance.',
    'Performed routine IPC checks and recorded in-process results in SAP to support batch traceability and data integrity.',
    'Supported investigations, deviations, CAPA and change control activities in TrackWise and contributed to audit / inspection readiness.',
    'Supported process stabilisation following equipment updates by adjusting compression parameters within approved ranges and escalating trends / deviations as required.',
]
for b in abbvie_bullets:
    add_bullet(doc, b)

add_body(doc, 'Regeneron, Raheen, Limerick - Biotechnician Specialist Nov 2021 - Mar 2022', bold=True)
for b in [
    'Supported GMP downstream purification operations including chromatography (HIC, AEX / CEX) and virus filtration; assisted with column equilibration, loading, washing, elution and strip activities per SOPs.',
    'Executed CIP / SIP cycles for skids, tanks and piping; performed line clearances and equipment cleaning with GDP / ALCOA+ documentation.',
    'Performed in-process sampling and basic testing (for example pH and conductivity) and recorded results contemporaneously to ALCOA+ standards.',
]:
    add_bullet(doc, b)

add_body(doc, 'Janssen Pharmaceutical, Cork - Quality Control Analyst Jan 2020 - Jun 2021', bold=True)
for b in [
    'Performed routine physicochemical QC testing under cGMP, including Karl Fischer water content, pH, conductivity, FTIR / IR identity testing, loss on drying, residue on ignition / sulphated ash, density / specific gravity, refractive index, acid / base titrations, insoluble matter and appearance.',
    'Acted as Karl Fischer SME / trainer, performing routine and non-routine testing, troubleshooting and instrument calibration / verification while training analysts and supporting OOS / OOT investigations to closure.',
    'Logged, labelled and tracked samples in LIMS, maintaining chain of custody and ALCOA+ data integrity from sample receipt through reporting.',
    'Investigated OOS / OOT results with QA, documented investigations in TrackWise and supported CAPA / change control actions.',
    'Served as 5S representative, redesigning lab storage and introducing Kanban controls to improve reagent availability and reduce search time.',
]:
    add_bullet(doc, b)

add_heading(doc, 'EDUCATION')
add_body(doc, 'ATU Sligo - Postgraduate Diploma (NFQ Level 9), Commissioning, Qualification and Validation (CQV) Completed May 2025 | First Class Honours (1.1)', bold=True)
add_body(doc, 'Coursework included URS / DQ, FAT / SAT, IQ / OQ / PQ, risk assessment (FMEA), CSV, P&IDs, facility design / cleanrooms and tech transfer / validation.')
add_body(doc, 'University College Cork - BSc Biotechnology (2:1) 2015 - 2019', bold=True)
add_body(doc, 'Dissertation: Recombinant transaminase expression (E. coli).')

doc.save('Mark_Healy_CV.docx')
print('Created Mark_Healy_CV.docx')
